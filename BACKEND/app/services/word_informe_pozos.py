"""
Generación del Informe General de Pozos a Tierra en formato Word (.docx).
Réplica exacta del sistema heredado.
"""
from __future__ import annotations

import io
import logging
import os
from datetime import date
from typing import Optional

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from PIL import Image as PILImage

from .word_helpers import aplicar_encabezado_informe_general

logger = logging.getLogger(__name__)

_MESES_ES = [
    "", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]

# Body table full width (cm)
_BW = 17.6


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cm_to_twips(cm: float) -> int:
    """1 cm = 567 twips (1440 twips/in ÷ 2.54 cm/in)."""
    return int(cm * 567)


def _set_col_width(cell, width_cm: float) -> None:
    """Set exact cell width in cm, removing any pre-existing <w:tcW>."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):   # remove duplicates Word would ignore
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_tbl_grid(table, col_widths_cm: list[float]) -> None:
    """Inject a proper <w:tblGrid> so Word knows the column layout."""
    tbl = table._tbl
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(_cm_to_twips(w)))
        tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _no_borders(table) -> None:
    """Remove all visible borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "none")
        bd.set(qn("w:sz"), "0")
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), "auto")
        tblBorders.append(bd)
    tblPr.append(tblBorders)



def _single_borders(table, color="000000", sz=6) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), str(sz))
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tblBorders.append(bd)
    tblPr.append(tblBorders)


def _set_table_width(table, width_cm: float) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _run(para, text: str, bold=False, size_pt: float = 11,
         color: Optional[RGBColor] = None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    if color:
        r.font.color.rgb = color
    return r


def _para_heading(doc: Document, text: str, size_pt: float, bold=True, bookmark: str = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, text, bold=bold, size_pt=size_pt)
    if bookmark:
        _add_bookmark(p, bookmark)


def _para_body(doc: Document, text: str, justify=True) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, size_pt=11)


def _para_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    _run(p, "• " + text, size_pt=11)


def _gray_header_row(row, texts: list[str], size_pt=11) -> None:
    """Fill a table row with gray background and bold text."""
    for cell, txt in zip(row.cells, texts):
        _set_cell_bg(cell, "F2F2F2")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, bold=True, size_pt=size_pt)


# ─────────────────────────────────────────────────────────────────────────────
# Bookmark / hyperlink helpers (índice interactivo)
# ─────────────────────────────────────────────────────────────────────────────

_bm_id_counter = [0]  # mutable so it resets per document call


def _next_bm_id() -> int:
    bid = _bm_id_counter[0]
    _bm_id_counter[0] += 1
    return bid


def _add_bookmark(paragraph, bookmark_name: str) -> None:
    bid = str(_next_bm_id())
    p = paragraph._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), bid)
    bm_start.set(qn("w:name"), bookmark_name)
    p.insert(0, bm_start)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), bid)
    p.append(bm_end)


def _add_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


# ─────────────────────────────────────────────────────────────────────────────
# Image helpers
# ─────────────────────────────────────────────────────────────────────────────

def _url_to_jpeg(url: str, timeout=15) -> Optional[io.BytesIO]:
    try:
        resp = requests.get(url, timeout=timeout)
        if resp.status_code != 200 or not resp.content:
            return None
        img = PILImage.open(io.BytesIO(resp.content))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGBA")
            bg = PILImage.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        else:
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        return buf
    except Exception as exc:
        logger.warning("[word_pozos] imagen URL no descargada %s: %s", url, exc)
        return None


def _add_logo(para, path: str, width_cm: float) -> None:
    try:
        para.add_run().add_picture(path, width=Cm(width_cm))
    except Exception:
        para.add_run(path.split("/")[-1])


# ─────────────────────────────────────────────────────────────────────────────
# 1. PORTADA
# ─────────────────────────────────────────────────────────────────────────────

def _build_portada(doc: Document, cfg: dict, oc: str, provincia: str, fecha_str: str) -> None:
    oc_label = oc or "SIN-OC"
    prov_label = (provincia or "").upper()

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    for text, size in [
        (f"INFORME OC - {oc_label}", 18),
        (cfg.get("titulo_documento", "MANTENIMIENTO PREVENTIVO"), 16),
        (f"TALMA - {prov_label}", 16),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, bold=True, size_pt=size)

    doc.add_paragraph()
    doc.add_paragraph()

    today = date.today()
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_fecha, f"{_MESES_ES[today.month]} – {today.year}", bold=True, size_pt=14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Execution table
    tbl = doc.add_table(rows=2, cols=3)
    _single_borders(tbl)
    _set_table_width(tbl, _BW)
    _set_tbl_grid(tbl, [1.8, 8.8, 7.0])
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(2):
        _set_col_width(tbl.cell(ri, 0), 1.8)
        _set_col_width(tbl.cell(ri, 1), 8.8)
        _set_col_width(tbl.cell(ri, 2), 7.0)

    _gray_header_row(tbl.rows[0], ["N°", "DESCRIPCIÓN", "FECHA"])
    row1 = tbl.rows[1]
    for cell, txt in zip(row1.cells, ["01", "Fecha de ejecución de trabajo", "          /          /          "]):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, txt, size_pt=11)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# 3. ÍNDICE
# ─────────────────────────────────────────────────────────────────────────────

_INDENT_MAP = {0: Cm(0), 1: Cm(0.5), 2: Cm(1.0)}

_LEVEL0_BOOKMARKS_BASE = {
    "01. GENERALIDADES":                                         "sec_01",
    "02. ALCANCE":                                               "sec_02",
    "03. MARCO NORMATIVO":                                       "sec_03",
    "04. OBJETIVOS":                                             "sec_04",
    "05. PROCEDIMIENTO SST IMPLÍCITOS":                          "sec_05",
    "06. PERSONAL COMPROMETIDO":                                 "sec_06",
    "07. LISTA DE MATERIALES, HERRAMIENTAS Y EQUIPOS EMPLEADOS": "sec_07",
    "08. PROCEDIMIENTOS DE EJECUCIÓN":                           "sec_08",
    "09. EVIDENCIA FOTOGRÁFICA":                                 "sec_09",
    "10. CONCLUSIONES":                                          "sec_10",
    "11. OBSERVACIONES":                                         "sec_11",
    "12. RECOMENDACIONES":                                       "sec_12",
    "13. ANEXOS":                                                "sec_13",
}


def _indice_items(cfg: dict) -> list[tuple[int, str]]:
    norms = cfg.get("normativas", {})
    items = [
        (0, "01. GENERALIDADES"),
        (0, "02. ALCANCE"),
        (0, "03. MARCO NORMATIVO"),
        (1, "03.1. NORMATIVA NACIONAL"),
        *[(2, n["titulo"]) for n in norms.get("nacional", [])],
        (1, "03.2. NORMATIVA INTERNACIONAL"),
        *[(2, n["titulo"]) for n in norms.get("internacional", [])],
        (0, "04. OBJETIVOS"),
        (0, "05. PROCEDIMIENTO SST IMPLÍCITOS"),
        (1, "05.1. USO DEL EPP DE ACUERDO AL TIPO DE TRABAJO"),
        (1, "05.2. PROCEDIMIENTO DE BLOQUEO Y ETIQUETADO LOTO"),
        (0, "06. PERSONAL COMPROMETIDO"),
        (0, "07. LISTA DE MATERIALES, HERRAMIENTAS Y EQUIPOS EMPLEADOS"),
        (0, "08. PROCEDIMIENTOS DE EJECUCIÓN"),
        (1, "08.1. PREPARATIVOS PREVIOS A LA EJECUCIÓN"),
        (1, "08.2. EJECUCIÓN DE TRABAJO"),
        (0, "09. EVIDENCIA FOTOGRÁFICA"),
        (0, "10. CONCLUSIONES"),
        (0, "11. OBSERVACIONES"),
        (0, "12. RECOMENDACIONES"),
    ]
    if cfg.get("tiene_anexos"):
        items.append((0, "13. ANEXOS"))
    return items


def _build_indice(doc: Document, cfg: dict) -> None:
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    _run(p_title, "Contenido", bold=True, size_pt=13)

    for level, text in _indice_items(cfg):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = _INDENT_MAP[level]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        bm = _LEVEL0_BOOKMARKS_BASE.get(text) if level == 0 else None
        if bm:
            _add_hyperlink(p, text, bm)
        else:
            _run(p, text, size_pt=10)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Static section texts
# ─────────────────────────────────────────────────────────────────────────────

def _sec_generalidades(doc: Document, cfg: dict, provincia: str) -> None:
    prov = (provincia or "").upper()
    for para in cfg.get("generalidades", []):
        _para_body(doc, para.format(prov=prov))


def _sec_alcance(doc: Document, cfg: dict, provincia: str) -> None:
    prov = (provincia or "").upper()
    for para in cfg.get("alcance", []):
        _para_body(doc, para.format(prov=prov))


def _sec_marco_normativo(doc: Document, cfg: dict) -> None:
    norms = cfg.get("normativas", {})
    _para_heading(doc, "03.1. NORMATIVA NACIONAL", 12)
    for norm in norms.get("nacional", []):
        _para_heading(doc, norm["titulo"], 11)
        for para in norm["cuerpo"]:
            _para_body(doc, para)
    _para_heading(doc, "03.2. NORMATIVA INTERNACIONAL", 12)
    for norm in norms.get("internacional", []):
        _para_heading(doc, norm["titulo"], 11)
        for para in norm["cuerpo"]:
            _para_body(doc, para)


def _sec_objetivos(doc: Document, cfg: dict) -> None:
    for b in cfg.get("objetivos", []):
        doc.add_paragraph(b, style="List Bullet")


def _sec_epp_intro(doc: Document, cfg: dict) -> None:
    _para_body(doc, cfg.get("epp_intro", ""))


def _sec_loto(doc: Document, cfg: dict) -> None:
    _para_body(doc, cfg.get("loto_intro", ""))
    for b in cfg.get("loto_bullets", []):
        doc.add_paragraph(b, style="List Bullet")


def _sec_personal_intro(doc: Document, cfg: dict) -> None:
    for para in cfg.get("personal_intro", []):
        _para_body(doc, para)


def _sec_preparativos(doc: Document, cfg: dict) -> None:
    for b in cfg.get("preparativos", []):
        doc.add_paragraph(b, style="List Bullet")


def _sec_ejecucion(doc: Document, cfg: dict) -> None:
    for b in cfg.get("ejecucion", []):
        doc.add_paragraph(b, style="List Bullet")


def _sec_conclusiones(doc: Document, cfg: dict) -> None:
    for b in cfg.get("conclusiones", []):
        doc.add_paragraph(b, style="List Bullet")


def _sec_observaciones(doc: Document, cfg: dict) -> None:
    _para_body(doc, cfg.get("observaciones", ""))


def _sec_recomendaciones(doc: Document, cfg: dict) -> None:
    _para_body(doc, cfg.get("recomendaciones", ""))


def _sec_anexos(doc: Document, cfg: dict) -> None:
    for item in cfg.get("anexos", []):
        _para_body(doc, item)


# ─────────────────────────────────────────────────────────────────────────────
# Dynamic tables
# ─────────────────────────────────────────────────────────────────────────────

def _epp_es_necesario(nombre: str) -> str:
    n = nombre.upper()
    if "CARETA" in n or "ARNES" in n or "ARNÉS" in n:
        return "NO"
    return "SI – Plataforma SST"


def _table_epp(doc: Document, epps: list[str]) -> None:
    """2 cols: EPP | ES NECESARIO — 8.8+8.8 cm."""
    tbl = doc.add_table(rows=1 + len(epps), cols=2)
    tbl.style = "Table Grid"
    _single_borders(tbl)
    _set_table_width(tbl, _BW)
    _set_tbl_grid(tbl, [8.8, 8.8])
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(1 + len(epps)):
        _set_col_width(tbl.cell(ri, 0), 8.8)
        _set_col_width(tbl.cell(ri, 1), 8.8)

    _gray_header_row(tbl.rows[0], ["EPP", "ES NECESARIO"])

    for i, epp in enumerate(epps, start=1):
        row = tbl.rows[i]
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _run(row.cells[0].paragraphs[0], epp, size_pt=11)
        _run(row.cells[1].paragraphs[0], _epp_es_necesario(epp), size_pt=11)


def _cargo_resp(rol: str) -> tuple[str, str]:
    if any(k in rol.lower() for k in ("supervis", "jefe", "lider", "líder")):
        return "Coordinador / Supervisor", "Supervisión de trabajos"
    return "Técnico ejecutor", "Ejecución de trabajos"


def _table_personal(doc: Document, personal: list[dict]) -> None:
    """4 cols: APELLIDOS | NOMBRE | CARGO | RESPONSABILIDADES — 4.4cm each."""
    tbl = doc.add_table(rows=1 + len(personal), cols=4)
    tbl.style = "Table Grid"
    _single_borders(tbl)
    _set_table_width(tbl, _BW)
    _set_tbl_grid(tbl, [4.4, 4.4, 4.4, 4.4])
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(1 + len(personal)):
        for ci in range(4):
            _set_col_width(tbl.cell(ri, ci), 4.4)

    _gray_header_row(tbl.rows[0], ["APELLIDOS", "NOMBRE", "CARGO", "RESPONSABILIDADES"])

    for i, pers in enumerate(personal, start=1):
        nombre_full = pers.get("nombre", "").strip()
        parts = nombre_full.split(" ", 1)
        nombre   = parts[0].upper() if parts else ""
        apellido = (parts[1].upper() + ",") if len(parts) == 2 else ""
        cargo, resp = _cargo_resp(pers.get("rol", "Técnico"))
        row = tbl.rows[i]
        for ci, txt in enumerate([apellido, nombre, cargo, resp]):
            row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _run(row.cells[ci].paragraphs[0], txt, size_pt=11)


def _table_materiales(doc: Document, materiales: list[dict]) -> None:
    """3 cols: MATERIALES | UNIDAD | CARACTERÍSTICAS — 7.1+3.5+7.0 cm."""
    tbl = doc.add_table(rows=1 + len(materiales), cols=3)
    tbl.style = "Table Grid"
    _single_borders(tbl)
    _set_table_width(tbl, _BW)
    _set_tbl_grid(tbl, [7.1, 3.5, 7.0])
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(1 + len(materiales)):
        for ci, w in enumerate([7.1, 3.5, 7.0]):
            _set_col_width(tbl.cell(ri, ci), w)

    _gray_header_row(tbl.rows[0], ["MATERIALES", "UNIDAD", "CARACTERÍSTICAS"])

    for i, m in enumerate(materiales, start=1):
        row = tbl.rows[i]
        for ci, txt in enumerate([
            m.get("nombre", ""),
            m.get("unidad", "Und."),
            m.get("caracteristicas", "-"),
        ]):
            row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _run(row.cells[ci].paragraphs[0], txt, size_pt=11)


def _table_herramientas(doc: Document, herramientas: list[dict]) -> None:
    """3 cols: HERRAMIENTAS Y EQUIPOS | UNIDAD | CARACTERÍSTICAS — 7.1+3.5+7.0 cm."""
    tbl = doc.add_table(rows=1 + len(herramientas), cols=3)
    tbl.style = "Table Grid"
    _single_borders(tbl)
    _set_table_width(tbl, _BW)
    _set_tbl_grid(tbl, [7.1, 3.5, 7.0])
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(1 + len(herramientas)):
        for ci, w in enumerate([7.1, 3.5, 7.0]):
            _set_col_width(tbl.cell(ri, ci), w)

    _gray_header_row(tbl.rows[0], ["HERRAMIENTAS Y EQUIPOS", "UNIDAD", "CARACTERÍSTICAS"])

    for i, h in enumerate(herramientas, start=1):
        row = tbl.rows[i]
        nombre = h.get("nombre", "")
        if h.get("serie") and h["serie"] != "-":
            nombre = f"{nombre} {h['serie']}".strip()
        for ci, txt in enumerate([nombre, "1 Ud.", h.get("marca", "-")]):
            row.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _run(row.cells[ci].paragraphs[0], txt, size_pt=11)


# ─────────────────────────────────────────────────────────────────────────────
# Section 09: Evidence photos — 2x2 grid
# ─────────────────────────────────────────────────────────────────────────────

def _evidence_grid(doc: Document, evidencias_por_equipo: list[dict]) -> None:
    """
    For each equipment group: title + 2-column tables (image row / caption row).
    Odd photo at the end fills left column; right column stays blank.
    Each image is capped at Cm(7.5) so the table never overflows A4 width.
    """
    W_COL = 7.5   # cm — each of the two columns (total 15 cm, within 17.6 cm body)

    for grupo in evidencias_por_equipo:
        nombre = grupo.get("nombre", "Equipo")
        fotos  = grupo.get("fotos", [])

        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p_t, f"EJECUCIÓN DE MANTENIMIENTO PREVENTIVO – {nombre.upper()}",
             bold=True, size_pt=11)

        if not fotos:
            _para_body(doc, "(Sin evidencias fotográficas registradas)", justify=False)
            doc.add_paragraph()
            continue

        i = 0
        while i < len(fotos):
            foto_izq = fotos[i]
            foto_der = fotos[i + 1] if i + 1 < len(fotos) else None

            # Always a 2-column / 2-row table: row 0 = images, row 1 = captions
            tbl = doc.add_table(rows=2, cols=2)
            tbl.style = "Table Grid"
            _set_table_width(tbl, W_COL * 2)
            _set_tbl_grid(tbl, [W_COL, W_COL])
            for ri in range(2):
                for ci in range(2):
                    _set_col_width(tbl.cell(ri, ci), W_COL)

            # ── Row 0: images ──────────────────────────────────────────────
            for ci, foto in enumerate([foto_izq, foto_der]):
                c_img = tbl.cell(0, ci)
                c_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p_img = c_img.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if foto is None:
                    continue   # leave right cell blank when photo count is odd
                buf = _url_to_jpeg(foto["url"])
                if buf:
                    p_img.add_run().add_picture(buf, width=Cm(W_COL))
                else:
                    _run(p_img, "[imagen no disponible]", size_pt=9)

            # ── Row 1: captions ────────────────────────────────────────────
            for ci, foto in enumerate([foto_izq, foto_der]):
                c_obs = tbl.cell(1, ci)
                p_obs = c_obs.paragraphs[0]
                p_obs.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_obs.paragraph_format.space_before = Pt(2)
                if foto is None:
                    continue
                _run(p_obs, foto.get("obs", ""), size_pt=9)

            doc.add_paragraph()
            i += 2

        doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def generar_word_pozos(
    cfg: dict,
    oc: str,
    ubicacion: str,
    equipos: list[dict],
    epps: list[str],
    herramientas: list[dict],
    materiales: list[dict],
    personal: list[dict],
    evidencias_por_equipo: list[dict],
    razon_social: str = "",
) -> bytes:
    """
    Genera el Word del Informe General y devuelve los bytes.
    cfg: diccionario de contenido del equipo (de report_templates.CONFIGS).
    """
    doc = Document()

    # Page margins to match legacy (usable width ≈ 17.6 cm)
    sec = doc.sections[0]
    sec.left_margin  = Cm(1.7)
    sec.right_margin = Cm(1.7)
    sec.top_margin   = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # Global font style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    today = date.today()
    fecha_str = f"{today.day:02d}/{today.month:02d}/{today.year}"

    # Reset bookmark ID counter for this document
    _bm_id_counter[0] = 0

    # ── 1. Header ──────────────────────────────────────────────────────────────
    aplicar_encabezado_informe_general(doc, {
        "oc":           oc,
        "tipo_equipo":  cfg.get("tipo_label", ""),
        "lugar":        ubicacion,
        "razon_social": razon_social,
    })

    # ── 2. Portada ─────────────────────────────────────────────────────────────
    _build_portada(doc, cfg, oc, ubicacion, fecha_str)

    # ── 3. Índice ──────────────────────────────────────────────────────────────
    _build_indice(doc, cfg)

    # ── 4-12. Secciones de contenido ──────────────────────────────────────────

    # 01. GENERALIDADES
    _para_heading(doc, "01. GENERALIDADES", 14, bookmark="sec_01")
    _sec_generalidades(doc, cfg, ubicacion)
    doc.add_paragraph()

    # 02. ALCANCE
    _para_heading(doc, "02. ALCANCE", 14, bookmark="sec_02")
    _sec_alcance(doc, cfg, ubicacion)
    doc.add_paragraph()

    # 03. MARCO NORMATIVO
    _para_heading(doc, "03. MARCO NORMATIVO", 14, bookmark="sec_03")
    _sec_marco_normativo(doc, cfg)
    doc.add_paragraph()

    # 04. OBJETIVOS
    _para_heading(doc, "04. OBJETIVOS", 14, bookmark="sec_04")
    _sec_objetivos(doc, cfg)
    doc.add_paragraph()

    # 05. PROCEDIMIENTO SST
    _para_heading(doc, "05. PROCEDIMIENTO SST IMPLÍCITOS", 14, bookmark="sec_05")
    _para_heading(doc, "05.1. USO DEL EPP DE ACUERDO AL TIPO DE TRABAJO", 12)
    _sec_epp_intro(doc, cfg)
    _table_epp(doc, epps)
    doc.add_paragraph()

    _para_heading(doc, "05.2. PROCEDIMIENTO DE BLOQUEO Y ETIQUETADO LOTO", 12)
    _sec_loto(doc, cfg)
    doc.add_paragraph()

    # 06. PERSONAL COMPROMETIDO
    _para_heading(doc, "06. PERSONAL COMPROMETIDO", 14, bookmark="sec_06")
    _sec_personal_intro(doc, cfg)
    _table_personal(doc, personal)
    doc.add_paragraph()

    # 07. MATERIALES, HERRAMIENTAS Y EQUIPOS
    _para_heading(doc, "07. LISTA DE MATERIALES, HERRAMIENTAS Y EQUIPOS EMPLEADOS", 14, bookmark="sec_07")
    _table_materiales(doc, materiales)
    doc.add_paragraph()
    _table_herramientas(doc, herramientas)
    doc.add_paragraph()

    # 08. PROCEDIMIENTOS DE EJECUCIÓN
    _para_heading(doc, "08. PROCEDIMIENTOS DE EJECUCIÓN", 14, bookmark="sec_08")
    _para_heading(doc, "08.1. PREPARATIVOS PREVIOS A LA EJECUCIÓN", 12)
    _sec_preparativos(doc, cfg)
    doc.add_paragraph()

    _para_heading(doc, "08.2. EJECUCIÓN DE TRABAJO", 12)
    _sec_ejecucion(doc, cfg)
    doc.add_paragraph()

    # 09. EVIDENCIA FOTOGRÁFICA
    _para_heading(doc, "09. EVIDENCIA FOTOGRÁFICA", 14, bookmark="sec_09")
    _evidence_grid(doc, evidencias_por_equipo)

    # 10. CONCLUSIONES
    _para_heading(doc, "10. CONCLUSIONES", 14, bookmark="sec_10")
    _sec_conclusiones(doc, cfg)
    doc.add_paragraph()

    # 11. OBSERVACIONES
    _para_heading(doc, "11. OBSERVACIONES", 14, bookmark="sec_11")
    _sec_observaciones(doc, cfg)
    doc.add_paragraph()

    # 12. RECOMENDACIONES
    _para_heading(doc, "12. RECOMENDACIONES", 14, bookmark="sec_12")
    _sec_recomendaciones(doc, cfg)

    # 13. ANEXOS (solo si el tipo lo requiere)
    if cfg.get("tiene_anexos"):
        doc.add_paragraph()
        _para_heading(doc, "13. ANEXOS", 14, bookmark="sec_13")
        _sec_anexos(doc, cfg)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

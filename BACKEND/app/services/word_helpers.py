"""
Word document header helpers — dos implementaciones completamente independientes:
  - aplicar_encabezado_garantia       : imágenes flotantes (Carta de Garantía)
  - aplicar_encabezado_informe_general: tabla 3×4 con bordes grises (Informes Generales)
"""
from __future__ import annotations

import os
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_LOGO_HEADER   = os.path.abspath(os.path.join("assets", "headerLogo.png"))
_DESIGN_HEADER = os.path.abspath(os.path.join("assets", "headerDesign.png"))

_WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
_A_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
_W_NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Anchos de columna del encabezado tabla (cm) — 50 / 25 / 25 del ancho útil 17.6 cm
_H_COL0 = 8.80   # 50 %
_H_COL1 = 4.40   # 25 %
_H_COL2 = 4.40   # 25 %


# ─── Helpers internos ────────────────────────────────────────────────────────

def _add_floating_picture(paragraph, image_path, width, pos_x, pos_y) -> None:
    """Inserta imagen flotante anclada en coordenadas absolutas de página."""
    run = paragraph.add_run()
    run.add_picture(image_path, width=width)

    drawing = run._r.find(f'.//{{{_W_NS}}}drawing')
    inline  = drawing.find(f'{{{_WP_NS}}}inline')

    extent  = inline.find(f'{{{_WP_NS}}}extent')
    eff_ext = inline.find(f'{{{_WP_NS}}}effectExtent')
    doc_pr  = inline.find(f'{{{_WP_NS}}}docPr')
    cnv_pr  = inline.find(f'{{{_WP_NS}}}cNvGraphicFramePr')
    graphic = inline.find(f'{{{_A_NS}}}graphic')

    anchor_xml = (
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"'
        f' relativeHeight="251658240" behindDoc="1" locked="0"'
        f' layoutInCell="1" allowOverlap="1"'
        f' xmlns:wp="{_WP_NS}">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{int(pos_x)}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{int(pos_y)}</wp:posOffset></wp:positionV>'
        f'<wp:wrapNone/>'
        f'</wp:anchor>'
    )
    anchor = parse_xml(anchor_xml)

    wrap_pos = list(anchor).index(anchor.find(f'{{{_WP_NS}}}wrapNone'))
    for elem in [e for e in (extent, eff_ext) if e is not None]:
        anchor.insert(wrap_pos, elem)
        wrap_pos += 1

    for elem in [e for e in (doc_pr, cnv_pr, graphic) if e is not None]:
        anchor.append(elem)

    drawing.replace(inline, anchor)


def _cm_to_twips(cm: float) -> int:
    return int(cm * 567)


def _set_col_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_tbl_grid(table, col_widths_cm: list) -> None:
    tbl = table._tbl
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(_cm_to_twips(w)))
        tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _set_table_width(table, width_cm: float) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _apply_table_borders(table, color: str = "BFBFBF", sz: int = 4) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), str(sz))
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tblBorders.append(bd)
    tblPr.append(tblBorders)


def _cell_write(cell, text: str, bold: bool = False, size_pt: float = 8,
                color: Optional[RGBColor] = None) -> None:
    """Escribe texto en la primera celda de un párrafo con espaciado cero."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)
    r.font.name = "Arial"
    if color:
        r.font.color.rgb = color


def _cell_write_mixed(cell, label: str, value: str, size_pt: float = 8) -> None:
    """Escribe en la celda con 'label' en negrita seguido de 'value' normal."""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r_lbl = p.add_run(label)
    r_lbl.bold = True
    r_lbl.font.size = Pt(size_pt)
    r_lbl.font.name = "Arial"
    r_val = p.add_run(value)
    r_val.font.size = Pt(size_pt)
    r_val.font.name = "Arial"


# ─── Función pública 1: Carta de Garantía ────────────────────────────────────

def aplicar_encabezado_garantia(document: Document) -> None:
    """
    Encabezado exclusivo de la Carta de Garantía.
    Logo flotante izquierdo (X=1.27cm, Y=0.5cm) + banner flotante derecho (X=10cm, Y=0cm).
    Ambas imágenes se anclan en coordenadas absolutas de página (detrás del texto).
    """
    sec = document.sections[0]
    sec.header_distance = Cm(0.5)
    sec.top_margin = Cm(1.27)

    header = sec.header
    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    if os.path.exists(_LOGO_HEADER):
        _add_floating_picture(p, _LOGO_HEADER,   width=Cm(3),  pos_x=Cm(1.27), pos_y=Cm(0.5))
    if os.path.exists(_DESIGN_HEADER):
        _add_floating_picture(p, _DESIGN_HEADER, width=Cm(11), pos_x=Cm(10),   pos_y=Cm(0))


# ─── Función pública 2: Informes Generales ───────────────────────────────────

def aplicar_encabezado_informe_general(document: Document, datos: dict) -> None:
    """
    Encabezado exclusivo de los Informes Generales (Pozos, Tableros, Luminarias,
    Pararrayos, UPS). Estructura: tabla 3 columnas × 4 filas con bordes grises.

    datos = {
        "oc":           "OC-123",
        "tipo_equipo":  "POZOS",          # cfg["tipo_label"]
        "lugar":        "CALLAO",
        "razon_social": "TALMA",
    }
    """
    oc           = (datos.get("oc")           or "SIN-OC").upper()
    tipo_equipo  = (datos.get("tipo_equipo")  or "").upper()
    lugar        = (datos.get("lugar")        or "").upper()
    razon_social = (datos.get("razon_social") or "").upper()

    sec = document.sections[0]
    sec.header_distance = Cm(0.5)
    header = sec.header
    hdr_elem = header._element

    # Vaciar el header (el párrafo por defecto se reemplaza por la tabla)
    for child in list(hdr_elem):
        hdr_elem.remove(child)

    # ── Tabla 3 × 4 ──────────────────────────────────────────────────────────
    table = header.add_table(rows=4, cols=3)
    _set_table_width(table, _H_COL0 + _H_COL1 + _H_COL2)
    _set_tbl_grid(table, [_H_COL0, _H_COL1, _H_COL2])
    _apply_table_borders(table, color="BFBFBF", sz=4)

    GRAY = RGBColor(89, 89, 89)
    SZ   = 8

    # ── Fila 0: Título del informe | Logo (merge cols 1+2) ────────────────────
    row0 = table.rows[0]
    c00, c01, c02 = row0.cells
    _set_col_width(c00, _H_COL0)
    _set_col_width(c01, _H_COL1)
    _set_col_width(c02, _H_COL2)

    c00.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    titulo = f"INFORME OC - {oc} MANTENIMIENTO PREVENTIVO {tipo_equipo} {lugar}".strip()
    _cell_write(c00, titulo, bold=True, size_pt=SZ)

    merged_logo = c01.merge(c02)
    merged_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = merged_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after  = Pt(0)
    if os.path.exists(_LOGO_HEADER):
        p_logo.add_run().add_picture(_LOGO_HEADER, width=Cm(3.5))
    else:
        p_logo.add_run("LOGO")

    # ── Fila 1: Ubicación | Elaborado por | Oficina ───────────────────────────
    row1 = table.rows[1]
    c10, c11, c12 = row1.cells
    for cell, w in zip((c10, c11, c12), (_H_COL0, _H_COL1, _H_COL2)):
        _set_col_width(cell, w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _cell_write_mixed(c10, "UBICACIÓN: ", f"{razon_social} - {lugar}", size_pt=SZ)
    _cell_write(c11, "Elaborado por:",     size_pt=SZ, color=GRAY)
    _cell_write(c12, "Oficina de proyectos", size_pt=SZ, color=GRAY)

    # ── Fila 2: Especialidad | Revisado por | Oficina ─────────────────────────
    row2 = table.rows[2]
    c20, c21, c22 = row2.cells
    for cell, w in zip((c20, c21, c22), (_H_COL0, _H_COL1, _H_COL2)):
        _set_col_width(cell, w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _cell_write_mixed(c20, "ESPECIALIDAD: ", "INSTALACIONES ELÉCTRICAS", size_pt=SZ)
    _cell_write(c21, "Revisado por:",       size_pt=SZ, color=GRAY)
    _cell_write(c22, "Oficina de proyectos", size_pt=SZ, color=GRAY)

    # ── Fila 3: (vacía) | Aprobado por | Oficina ─────────────────────────────
    row3 = table.rows[3]
    c30, c31, c32 = row3.cells
    for cell, w in zip((c30, c31, c32), (_H_COL0, _H_COL1, _H_COL2)):
        _set_col_width(cell, w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p30 = c30.paragraphs[0]
    p30.paragraph_format.space_before = Pt(0)
    p30.paragraph_format.space_after  = Pt(0)

    _cell_write(c31, "Aprobado por:",        size_pt=SZ, color=GRAY)
    _cell_write(c32, "Oficina de proyectos", size_pt=SZ, color=GRAY)

    # Párrafo final obligatorio (requerido por OOXML para w:hdr)
    p_trail = OxmlElement("w:p")
    hdr_elem.append(p_trail)

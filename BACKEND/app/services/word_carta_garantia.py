"""
Generación de la Carta de Garantía de Servicio (.docx) con python-docx.
Diseño de una sola página — márgenes, spacings y fuentes calibradas
contra el documento de referencia Carta_Garantia_87.docx.
"""
from __future__ import annotations

import base64
import io
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_LOGO_HEADER   = os.path.abspath(os.path.join("assets", "headerLogo.png"))
_DESIGN_HEADER = os.path.abspath(os.path.join("assets", "headerDesign.png"))

_MESES_ES = [
    "", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]

# Ancho del cuerpo disponible con márgenes left/right de 2 cm en A4
_BW = 17.0   # 21 cm - 2*2 cm = 17 cm


# ─── XML helpers ─────────────────────────────────────────────────────────────

def _cm_to_twips(cm: float) -> int:
    return int(cm * 567)


def _set_col_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_tbl_grid(table, col_widths_cm: list[float]) -> None:
    tbl = table._tbl
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(_cm_to_twips(w)))
        tblGrid.append(gc)
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def _set_table_width(table, width_cm: float) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _no_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "none")
        bd.set(qn("w:sz"), "0")
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), "auto")
        tblBorders.append(bd)
    tblPr.append(tblBorders)


def _single_borders(table, color="000000", sz=6) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), str(sz))
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), color)
        tblBorders.append(bd)
    tblPr.append(tblBorders)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _run(para, text: str, bold=False, size_pt: float = 10,
         color: Optional[RGBColor] = None, italic=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size_pt)
    r.font.name = "Calibri"
    if color:
        r.font.color.rgb = color
    return r


def _spacing(para, before_pt: float = 0, after_pt: float = 0,
             line_pt: Optional[float] = None) -> None:
    """Aplica space_before, space_after y line_spacing a un párrafo."""
    fmt = para.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after  = Pt(after_pt)
    if line_pt is not None:
        fmt.line_spacing = Pt(line_pt)


def _mes_anio(fecha_str: str) -> str:
    """'2026-06-07' → 'Junio del 2026'."""
    try:
        d = datetime.fromisoformat(fecha_str)
        return f"{_MESES_ES[d.month]} del {d.year}"
    except Exception:
        return fecha_str


# ─── Encabezado de página ─────────────────────────────────────────────────────

def _build_header(doc: Document) -> None:
    """Tabla sin bordes: headerLogo.png izquierda | headerDesign.png derecha."""
    sec = doc.sections[0]
    sec.header_distance = Cm(0.2)
    header = sec.header

    for p in header.paragraphs:
        p.clear()

    col_logo   = 4.5
    col_design = _BW - col_logo   # ≈ 12.5 cm

    tbl = header.add_table(rows=1, cols=2, width=Cm(_BW))
    _no_borders(tbl)
    _set_tbl_grid(tbl, [col_logo, col_design])
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Logo izquierda
    c_logo = tbl.cell(0, 0)
    _set_col_width(c_logo, col_logo)
    c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(_LOGO_HEADER):
        p_logo.add_run().add_picture(_LOGO_HEADER, width=Cm(3.2))
    else:
        _run(p_logo, "E-System Tic Perú S.A.C.", bold=True, size_pt=10)

    # Diseño derecha
    c_design = tbl.cell(0, 1)
    _set_col_width(c_design, col_design)
    c_design.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_design = c_design.paragraphs[0]
    p_design.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(_DESIGN_HEADER):
        p_design.add_run().add_picture(_DESIGN_HEADER, width=Cm(col_design))


# ─── Pie de página oficial ────────────────────────────────────────────────────

def _build_footer(doc: Document) -> None:
    """Pie de página con datos oficiales de E-System Tic Perú S.A.C."""
    sec = doc.sections[0]
    sec.footer_distance = Cm(0.5)
    footer = sec.footer

    for p in footer.paragraphs:
        p.clear()

    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_footer, before_pt=2, after_pt=0)

    # Línea separadora superior via bordes del párrafo
    pPr = p_footer._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "AAAAAA")
    pBdr.append(top)
    pPr.append(pBdr)

    _run(
        p_footer,
        "RUC: 20451506235  |  Av. Bocanegra MZ. G62.LT10 – Callao  |  "
        "Telf. 981-443-214  |  www.esystemtic.com",
        size_pt=8, color=RGBColor(0x55, 0x55, 0x55),
    )


# ─── Función principal ────────────────────────────────────────────────────────

def generar_carta_garantia(
    equipos:               list[dict],
    razon_social:          str,
    fecha_operatividad:    str,
    vigencia_garantia:     str,
    lugar:                 str,
    direccion:             str,
    firma_verificador_b64: str,
    firma_gerente_b64:     str,
) -> bytes:
    """
    Genera la Carta de Garantía de Servicio para tableros eléctricos.
    Devuelve bytes del .docx ajustados para caber en una sola página.
    """
    doc = Document()

    # ── Márgenes (calibrados contra referencia para 1 sola página) ──────────
    sec = doc.sections[0]
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(1.27)
    sec.bottom_margin = Cm(1.27)

    # ── Fuente global por defecto ────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    _build_header(doc)
    _build_footer(doc)

    lugar_upper = (lugar or "").upper()

    # ── TÍTULO ───────────────────────────────────────────────────────────────
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(
        p_titulo,
        f"GARANTÍA DE SERVICIO MANTENIMIENTO PREVENTIVO "
        f"DE TABLEROS ELÉCTRICOS EN {lugar_upper}",
        bold=True, size_pt=11,
    )
    _spacing(p_titulo, before_pt=0, after_pt=6)

    # ── RAZÓN SOCIAL ─────────────────────────────────────────────────────────
    p_rs = doc.add_paragraph()
    _run(p_rs, "Razón Social: ", bold=True, size_pt=10)
    _run(p_rs, razon_social or "—", size_pt=10)
    _spacing(p_rs, before_pt=0, after_pt=2)

    # ── DIRECCIÓN ────────────────────────────────────────────────────────────
    p_dir = doc.add_paragraph()
    _run(p_dir, "DIRECCIÓN: ", bold=True, size_pt=10)
    _run(p_dir, direccion or "—", size_pt=10)
    _spacing(p_dir, before_pt=0, after_pt=5)

    # ── PÁRRAFO CUERPO ───────────────────────────────────────────────────────
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(
        p_body,
        f"Se expide la presente carta de garantía correspondiente al servicio de Mantenimiento "
        f"Preventivo de Tableros Eléctricos en {lugar_upper}, por parte de la empresa E-System Tic "
        f"Perú S.A.C. con RUC 20451506235 con el representante legal Edward Epifanio Galindo Segovia.",
        size_pt=10,
    )
    _spacing(p_body, before_pt=0, after_pt=5)

    # ── RELACIÓN DE TABLEROS ─────────────────────────────────────────────────
    p_rel = doc.add_paragraph()
    _run(p_rel, "RELACIÓN DE TABLEROS ELÉCTRICOS INTERVENIDOS", bold=True, size_pt=10)
    _spacing(p_rel, before_pt=3, after_pt=4)

    tbl_eq = doc.add_table(rows=1 + len(equipos), cols=1)
    tbl_eq.style = "Table Grid"
    _set_table_width(tbl_eq, _BW)
    tbl_eq.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = tbl_eq.rows[0]
    _set_cell_bg(hdr_row.cells[0], "F2F2F2")
    hc = hdr_row.cells[0].paragraphs[0]
    hc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(hc, "TABLERO ELÉCTRICO", bold=True, size_pt=10)
    _spacing(hc, before_pt=1, after_pt=1)

    for i, eq in enumerate(equipos, start=1):
        p_eq = tbl_eq.cell(i, 0).paragraphs[0]
        p_eq.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(p_eq, eq.get("nombre", "—"), size_pt=10)
        _spacing(p_eq, before_pt=1, after_pt=1)

    # Párrafo separador compacto después de la tabla de equipos
    p_gap1 = doc.add_paragraph()
    _spacing(p_gap1, before_pt=0, after_pt=4)

    # ── VIGENCIA ─────────────────────────────────────────────────────────────
    p_vig_label = doc.add_paragraph()
    _run(p_vig_label, "VIGENCIA DE GARANTÍA: GARANTÍA 12 MESES", bold=True, size_pt=10)
    _spacing(p_vig_label, before_pt=0, after_pt=4)

    tbl_vig = doc.add_table(rows=2, cols=2)
    _single_borders(tbl_vig)
    _set_table_width(tbl_vig, _BW)
    _set_tbl_grid(tbl_vig, [_BW / 2, _BW / 2])
    tbl_vig.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(2):
        _set_col_width(tbl_vig.cell(ri, 0), _BW / 2)
        _set_col_width(tbl_vig.cell(ri, 1), _BW / 2)

    _set_cell_bg(tbl_vig.cell(0, 0), "F2F2F2")
    _set_cell_bg(tbl_vig.cell(0, 1), "F2F2F2")
    for ci, label in enumerate(["Fecha de Mantenimiento", "Vigencia de garantía del servicio."]):
        c = tbl_vig.cell(0, ci)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, label, bold=True, size_pt=10)
        _spacing(p, before_pt=1, after_pt=1)

    for ci, fecha_str in enumerate([fecha_operatividad, vigencia_garantia]):
        c = tbl_vig.cell(1, ci)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, _mes_anio(fecha_str), size_pt=10)
        _spacing(p, before_pt=1, after_pt=1)

    # Párrafo separador compacto después de la tabla de vigencia
    p_gap2 = doc.add_paragraph()
    _spacing(p_gap2, before_pt=0, after_pt=5)

    # ── CIERRE ───────────────────────────────────────────────────────────────
    p_exp = doc.add_paragraph()
    p_exp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(
        p_exp,
        "Se expide la presente carta de garantía de servicio para los fines correspondientes.",
        size_pt=10,
    )
    _spacing(p_exp, before_pt=0, after_pt=4)

    # Párrafo separador antes de firmas
    p_gap3 = doc.add_paragraph()
    _spacing(p_gap3, before_pt=0, after_pt=20)

    # ── TABLA DE FIRMAS ───────────────────────────────────────────────────────
    tbl_firmas = doc.add_table(rows=1, cols=2)
    _no_borders(tbl_firmas)
    _set_table_width(tbl_firmas, _BW)
    _set_tbl_grid(tbl_firmas, [_BW / 2, _BW / 2])
    tbl_firmas.alignment = WD_TABLE_ALIGNMENT.CENTER

    _set_col_width(tbl_firmas.cell(0, 0), _BW / 2)
    _set_col_width(tbl_firmas.cell(0, 1), _BW / 2)

    def _fill_firma_cell(cell, firma_b64: str, lineas: list[str]) -> None:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for p in cell.paragraphs:
            p.clear()

        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p_img, before_pt=0, after_pt=1)

        if firma_b64:
            try:
                img_bytes = base64.b64decode(firma_b64)
                p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(3.5))
            except Exception:
                _run(p_img, "[imagen no disponible]", size_pt=8, italic=True)
        else:
            _run(p_img, " ", size_pt=10)

        for texto in lineas:
            p_l = cell.add_paragraph()
            p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p_l, texto, size_pt=10)
            _spacing(p_l, before_pt=0, after_pt=1)

    _fill_firma_cell(
        tbl_firmas.cell(0, 0),
        firma_verificador_b64,
        [
            "_________________________",
            "Personal Verificado",
            "Ing. Eléctrico y de Potencia",
            "Guerrero Sierra William Jesus",
            "CIP N° 218184",
        ],
    )

    _fill_firma_cell(
        tbl_firmas.cell(0, 1),
        firma_gerente_b64,
        [
            "_________________________",
            "Gerente General",
            "Ing. Galindo Segovia Edward.",
            "RUC: 20451506235",
        ],
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
